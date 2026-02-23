"""File analysis tool using vision models for image analysis."""

from .base import BaseTool, catalog_tool
from libs.messaging.redis import get_redis_client
from libs.llm import LLMMessage, MessageRole, get_provider
from libs.common.tool_catalog import get_tool_model_preference
from libs.common import get_logger
import sys
from datetime import UTC, datetime
from typing import Any

# Add parent directory to path for imports
sys.path.insert(0, "services/tool-workers")


logger = get_logger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@catalog_tool("analyze_file")
class AnalyzeFileTool(BaseTool):
    """Analyze uploaded files using vision models.

    Supports:
    - Images: JPEG, PNG, GIF, WebP (uses vision models)
    - PDFs: Text extraction and analysis (future: OCR)
    - Text files: Direct content analysis
    """

    async def execute(self, arguments: dict[str, Any], context: dict[str, Any]) -> str:
        """Execute file analysis using vision models.

        Flow:
        1. Validate file_id format (reject hallucinated IDs early)
        2. Check DB for cached analysis → return immediately if found
        3. File exists but not analyzed → try Redis for file data → analyze → cache
        4. File not in DB → clear error

        Args:
            arguments: Tool arguments (file_id, query)
            context: Execution context (job_id, tenant_id)

        Returns:
            Analysis result as string
        """
        import uuid as uuid_mod

        file_id = arguments["file_id"]
        query = arguments.get(
            "query", "Describe this file in full detail and don't miss any signle detail inside it")

        logger.info(
            "Analyzing file",
            file_id=file_id,
            query=query,
            job_id=context.get("job_id"),
        )

        # Step 1: Validate UUID format — catch hallucinated IDs like "image_id"
        try:
            uuid_mod.UUID(file_id)
        except ValueError:
            return (
                f"Error: Invalid file ID '{file_id}'. "
                "Expected a UUID (e.g., 'd0908dde-37e7-4441-bbc9-8d7f27a736c4'). "
                "Use the actual file_id from a previous upload or screenshot."
            )

        # Step 2: Check DB for existing cached analysis
        file_record = await self._get_file_record(file_id)

        if file_record and (file_record.extracted_text or file_record.analysis_description):
            # Prefer extracted_text (full content) over analysis_description (summary)
            cached = file_record.extracted_text or file_record.analysis_description
            logger.info(
                "Returning cached analysis from database",
                file_id=file_id,
                analyzed_at=str(file_record.analyzed_at),
                has_extracted_text=bool(file_record.extracted_text),
            )
            return cached

        # Step 3: No cached analysis — try to get file data from Redis
        try:
            redis = await get_redis_client()
            cache_key = f"file:{file_id}"
            cached_data = await redis.client.get(cache_key)

            if not cached_data:
                if file_record:
                    return (
                        f"File '{file_record.filename}' (type: {file_record.content_type}) "
                        "exists in the database but its data has expired from the cache. "
                        "The file has not been analyzed yet. Please ask the user to re-upload "
                        "the file so it can be analyzed."
                    )
                else:
                    return (
                        f"Error: File '{file_id}' not found. "
                        "No file with this ID exists in the database or cache. "
                        "Please verify the file_id is correct."
                    )

            # Parse file data
            import json
            import base64

            storage_payload = json.loads(cached_data)
            encoded_data = storage_payload["data"]
            metadata = storage_payload["metadata"]

            # Decode base64 data
            file_data_bytes = base64.b64decode(encoded_data)

            content_type = metadata.get("content_type", "")
            filename = metadata.get("filename", "unknown")

            logger.info(
                "File retrieved from Redis",
                file_id=file_id,
                filename=filename,
                content_type=content_type,
                size_bytes=len(file_data_bytes),
            )

            # Route based on content type
            extracted_text: str | None = None

            if content_type.startswith("image/"):
                result = await self._analyze_image(
                    file_data_bytes, content_type, query, filename
                )
            elif content_type == "application/pdf":
                result, extracted_text = await self._analyze_pdf(
                    file_data_bytes, query, filename, encoded_data
                )
            elif content_type == DOCX_MIME:
                result, extracted_text = await self._extract_docx(
                    file_data_bytes, filename
                )
            elif content_type == XLSX_MIME:
                result, extracted_text = await self._extract_xlsx(
                    file_data_bytes, filename
                )
            elif content_type.startswith("text/"):
                result, extracted_text = self._extract_text(
                    file_data_bytes, filename
                )
            else:
                return f"Error: Unsupported file type '{content_type}' for analysis."

            # Store analysis result in database for later retrieval
            await self._save_analysis_to_db(file_id, result, extracted_text)

            return result

        except Exception as e:
            logger.error(
                "File analysis failed",
                file_id=file_id,
                error=str(e),
                error_type=type(e).__name__,
            )
            return f"Error analyzing file: {str(e)}"

    async def _get_file_record(self, file_id: str):
        """Load the FileUpload record from the database.

        Args:
            file_id: UUID string of the file

        Returns:
            FileUpload model instance or None
        """
        try:
            import uuid as uuid_mod

            from libs.db.models import FileUpload
            from libs.db.session import get_session_context

            file_uuid = uuid_mod.UUID(file_id)

            async with get_session_context() as session:
                return await session.get(FileUpload, file_uuid)

        except Exception as e:
            logger.warning(
                "Failed to check database for file record",
                file_id=file_id,
                error=str(e),
            )
            return None

    async def _save_analysis_to_db(
        self,
        file_id: str,
        analysis: str,
        extracted_text: str | None = None,
    ) -> None:
        """Save analysis result to the file_uploads table for later retrieval.

        Args:
            file_id: UUID of the file
            analysis: The analysis text to store (description or raw text)
            extracted_text: Raw extracted text from document (PDF, DOCX, XLSX)
        """
        try:
            import uuid as uuid_mod

            from sqlalchemy import update

            from libs.db.models import FileUpload
            from libs.db.session import get_session_context

            file_uuid = uuid_mod.UUID(file_id)

            values: dict = {
                "analysis_description": analysis,
                "analyzed_at": datetime.now(UTC),
            }
            if extracted_text is not None:
                values["extracted_text"] = extracted_text

            async with get_session_context() as session:
                stmt = (
                    update(FileUpload)
                    .where(FileUpload.id == file_uuid)
                    .values(**values)
                )
                await session.execute(stmt)
                await session.commit()

            logger.info(
                "Analysis saved to database",
                file_id=file_id,
                analysis_length=len(analysis),
                has_extracted_text=extracted_text is not None,
            )
        except Exception as e:
            # Don't fail the tool if DB save fails — the analysis was still returned
            logger.error(
                "Failed to save analysis to database",
                file_id=file_id,
                error=str(e),
            )

    async def _analyze_image(
        self, file_data: bytes, content_type: str, query: str, filename: str
    ) -> str:
        """Analyze image using vision model with exhaustive detail.

        Args:
            file_data: Raw image bytes
            content_type: MIME type (e.g., image/jpeg)
            query: Analysis query from user
            filename: Original filename

        Returns:
            Analysis result
        """
        import base64

        # Re-encode to base64 for LLM
        encoded_data = base64.b64encode(file_data).decode("utf-8")

        # Use preferred provider/model from catalog
        provider_name, model_name = get_tool_model_preference("analyze_file")
        provider = get_provider(provider_name)

        # Build vision message with DETAILED description prompt
        message = LLMMessage(
            role=MessageRole.USER,
            content=[
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": content_type,
                        "data": encoded_data,
                    },
                },
                {
                    "type": "text",
                    "text": (
                        f"Analyze this image (filename: {filename}).\n\n"
                        f"Task: {query}\n\n"
                        "IMPORTANT: Provide an exhaustive, highly detailed description. "
                        "Do NOT summarize or abbreviate. Describe every visible element, "
                        "text, color, layout, structure, icon, label, and spatial relationship "
                        "in full detail. If there is text in the image, transcribe it exactly. "
                        "If this is a UI screenshot, describe every component, its state, "
                        "positioning, and content. Do not omit anything."
                    ),
                },
            ],
        )

        logger.info(
            "Calling vision model for image analysis",
            filename=filename,
            content_type=content_type,
            query_length=len(query),
        )

        response = await provider.complete(
            [message],
            model=model_name,
            max_tokens=4096,
        )

        logger.info(
            "Vision model analysis complete",
            filename=filename,
            output_tokens=response.output_tokens,
        )

        return response.content or "No analysis generated."

    async def _analyze_pdf(
        self,
        file_data: bytes,
        query: str,
        filename: str,
        encoded_data: str | None = None,
    ) -> tuple[str, str | None]:
        """Analyze PDF file.

        Tries text extraction first (cheap). If text is found, returns it
        directly — the main agent LLM will analyze the content in context.
        If no text is found, falls back to Anthropic's native PDF document
        block for scanned/image PDFs.

        Args:
            file_data: Raw PDF bytes
            query: Analysis query
            filename: Original filename
            encoded_data: Pre-encoded base64 string (avoids re-encoding)

        Returns:
            Tuple of (result_text, extracted_text_or_none)
        """
        try:
            import io
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_data))
            text_content = ""
            page_count = len(reader.pages)

            logger.info("Extracting text from PDF",
                        filename=filename, pages=page_count)

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\n--- Page {i+1} ---\n{page_text}"

            if not text_content.strip():
                # No embedded text — try native PDF analysis via Anthropic
                logger.info(
                    "No text extracted from PDF, falling back to native PDF analysis",
                    filename=filename,
                    pages=page_count,
                )
                result = await self._analyze_pdf_native(
                    file_data, query, filename, encoded_data
                )
                return result, None

            # Return extracted text directly — no secondary LLM call needed
            header = (
                f"Extracted text from '{filename}' "
                f"({page_count} page{'s' if page_count != 1 else ''}):\n"
            )
            result = header + text_content

            logger.info(
                "Returning extracted PDF text directly",
                filename=filename,
                content_length=len(text_content),
                pages=page_count,
            )

            return result, text_content

        except ImportError:
            return "Error: PDF analysis dependency (pypdf) is missing.", None
        except Exception as e:
            logger.error("PDF analysis failed", error=str(e))
            return f"Error analyzing PDF: {str(e)}", None

    async def _analyze_pdf_native(
        self,
        file_data: bytes,
        query: str,
        filename: str,
        encoded_data: str | None = None,
    ) -> str:
        """Analyze a scanned/image PDF using Anthropic's native document block.

        Args:
            file_data: Raw PDF bytes
            query: Analysis query
            filename: Original filename
            encoded_data: Pre-encoded base64 string (avoids re-encoding)

        Returns:
            Analysis result
        """
        import base64

        provider_name, model = get_tool_model_preference("analyze_file")

        if provider_name != "anthropic":
            return (
                f"Analysis Warning: No text could be extracted from '{filename}'. "
                "Native PDF analysis requires the Anthropic provider. "
                "Please convert the PDF to an image format (PNG/JPEG) for visual analysis."
            )

        if encoded_data is None:
            encoded_data = base64.b64encode(file_data).decode("utf-8")

        provider = get_provider(provider_name)

        message = LLMMessage(
            role=MessageRole.USER,
            content=[
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": encoded_data,
                    },
                },
                {
                    "type": "text",
                    "text": (
                        f"Analyze this PDF document (filename: {filename}).\n\n"
                        f"Task: {query}\n\n"
                        "IMPORTANT: Provide an exhaustive, highly detailed description. "
                        "Do NOT summarize or abbreviate. Describe every visible element, "
                        "text, table, figure, layout, and structure in full detail. "
                        "If there is text, transcribe it exactly."
                    ),
                },
            ],
        )

        logger.info(
            "Calling Anthropic with native PDF document block",
            filename=filename,
            model=model,
        )

        response = await provider.complete(
            [message],
            model=model,
            max_tokens=4096,
        )

        logger.info(
            "Native PDF analysis complete",
            filename=filename,
            output_tokens=response.output_tokens,
        )

        return response.content or "No analysis generated."

    def _extract_text(
        self, file_data: bytes, filename: str
    ) -> tuple[str, str | None]:
        """Extract text file content directly.

        Returns the raw text content — the main agent LLM will analyze it.

        Args:
            file_data: Raw text bytes
            filename: Original filename

        Returns:
            Tuple of (result_text, extracted_text_or_none)
        """
        try:
            text_content = file_data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text_content = file_data.decode("latin-1")
            except UnicodeDecodeError:
                return (
                    f"Error: Unable to decode text file '{filename}'. "
                    "File may be corrupted or in an unsupported encoding.",
                    None,
                )

        header = f"Extracted text from '{filename}':\n"
        result = header + text_content

        logger.info(
            "Returning extracted text file content directly",
            filename=filename,
            content_length=len(text_content),
        )

        return result, text_content

    async def _extract_docx(
        self, file_data: bytes, filename: str
    ) -> tuple[str, str | None]:
        """Extract text from Word (.docx) file.

        Extracts paragraphs (with heading/list style prefixes) and tables,
        returns the raw text directly for the main agent to analyze.

        Args:
            file_data: Raw .docx bytes
            filename: Original filename

        Returns:
            Tuple of (result_text, extracted_text_or_none)
        """
        try:
            import io

            from docx import Document
        except ImportError:
            return (
                "Error: Word document analysis dependency (python-docx) is missing.",
                None,
            )

        try:
            doc = Document(io.BytesIO(file_data))
            parts: list[str] = []

            # Extract paragraphs with style context
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower()
                if "heading" in style_name:
                    level = style_name.replace("heading", "").strip() or "1"
                    parts.append(f"{'#' * int(level)} {text}")
                elif "title" in style_name:
                    parts.append(f"# {text}")
                elif "list" in style_name:
                    parts.append(f"- {text}")
                else:
                    parts.append(text)

            # Extract tables
            for i, table in enumerate(doc.tables):
                parts.append(f"\n[Table {i + 1}]")
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append("| " + " | ".join(cells) + " |")
                    if row_idx == 0:
                        parts.append(
                            "| " + " | ".join("---" for _ in cells) + " |")

            text_content = "\n".join(parts)

            if not text_content.strip():
                return (
                    f"Analysis Warning: No content could be extracted from '{filename}'. "
                    "The Word document may be empty or contain only embedded objects.",
                    None,
                )

            header = f"Extracted text from '{filename}':\n"
            result = header + text_content

            logger.info(
                "Returning extracted Word document text directly",
                filename=filename,
                content_length=len(text_content),
            )

            return result, text_content

        except Exception as e:
            logger.error("Word document extraction failed", error=str(e))
            return f"Error extracting Word document: {str(e)}", None

    async def _extract_xlsx(
        self, file_data: bytes, filename: str
    ) -> tuple[str, str | None]:
        """Extract text from Excel (.xlsx) file.

        Extracts all sheets with rows formatted as pipe-separated values,
        returns the raw text directly for the main agent to analyze.

        Args:
            file_data: Raw .xlsx bytes
            filename: Original filename

        Returns:
            Tuple of (result_text, extracted_text_or_none)
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            return "Error: Excel analysis dependency (openpyxl) is missing.", None

        try:
            import io

            wb = load_workbook(
                io.BytesIO(file_data), read_only=True, data_only=True
            )
            parts: list[str] = []
            sheet_count = len(wb.sheetnames)

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"\n=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        parts.append("| " + " | ".join(cells) + " |")

            wb.close()

            text_content = "\n".join(parts)

            if not text_content.strip():
                return (
                    f"Analysis Warning: No data could be extracted from '{filename}'. "
                    "The Excel file may be empty.",
                    None,
                )
            header = (
                f"Extracted data from '{filename}' "
                f"({sheet_count} sheet{'s' if sheet_count != 1 else ''}):\n"
            )
            result = header + text_content

            logger.info(
                "Returning extracted Excel data directly",
                filename=filename,
                content_length=len(text_content),
                sheets=sheet_count,
            )

            return result, text_content

        except Exception as e:
            logger.error("Excel extraction failed", error=str(e))
            return f"Error extracting Excel file: {str(e)}", None
